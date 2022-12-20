import paramiko
import os
import gzip
import shutil
import pymysql
from pymysql.constants import CLIENT
from docxtpl import DocxTemplate
from datetime import datetime
from datetime import timedelta
import time
from datetime import timezone
from docx import Document
from docxcompose.composer import Composer

mydb = pymysql.connect(
    host="localhost",
    user="root",
    passwd="1111",
    database='UTM5',
    client_flag=CLIENT.MULTI_STATEMENTS)

files = []
lss = []
lstv = []
ids = 1
dts = datetime.today() + timedelta(days=1)
neopl = datetime.today() + timedelta(days=14)
s_data = datetime.today()
ost_data = datetime.today()
dt = datetime.today()
print(dt.strftime("%d.%m.%Y"))
print()
ts = int(dt.timestamp())
path = ['UTM5.users.sql', 'UTM5.accounts.sql', 'UTM5.service_links.sql', 'UTM5.periodic_services_data.sql',
        'UTM5.invoices.sql', 'UTM5.dtagg_periodic.sql', 'UTM5.discount_transactions_all.sql',
        'UTM5.payment_transactions.sql']
save = str(input('Обновить базу данных? Y/N: '))
if save == 'Y' or save == 'y':
    for paths in path:
        def sftp_get():
            local_path = 'D://Script/notifications/' + paths + '.gz'
            remote_path = '/share/utm5_backup/current/' + paths + '.gz'

            t = paramiko.Transport('172.16.100.29', 22)
            t.connect(username='admin', password='E,m.PfGfgbxf2045')
            sftp = paramiko.SFTPClient.from_transport(t)
            sftp.get(remote_path, local_path)
            t.close()


        if __name__ == '__main__':
            sftp_get()
        with gzip.open(paths + '.gz', 'rb') as f_in:
            with open(paths, 'wb') as f_out:
                shutil.copyfileobj(f_in, f_out)
        mycursor = mydb.cursor()
        x = '\n'.join(open(paths, 'r', encoding='UTF8').readlines())
        mycursor.execute(x)

print()
print('Интернет')
while ids <= 3000:
    mycursor = mydb.cursor()
    ls = 'SELECT basic_account FROM users WHERE is_deleted=0 AND NOT LOCATE ("office",login) AND NOT LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND id=%s'.format(
        ids)
    mycursor.execute(ls, ids)
    ls = mycursor.fetchall()
    ls = (str(ls)[2:-4])
    if ls != '':
        ls = int(ls)
        lss.append(ls)
    ids = ids + 1
for ls in lss:
    print(ls)
    bal = 'SELECT balance FROM accounts WHERE is_deleted=0 AND id=%s'.format(ls)
    mycursor.execute(bal, ls)
    bal = mycursor.fetchall()
    bal = round(float(str(bal)[2:-4]), 2)
    if bal < -10:
        payday =  'SELECT actual_date FROM payment_transactions_1234366079_1659083204 WHERE account_id=%s UNION ALL SELECT actual_date FROM payment_transactions WHERE account_id=%s'
        mycursor.execute(payday,(ls,ls))
        payday = mycursor.fetchall()
        if len(payday) == 0:
            continue
        else:
            payday = int(str(payday[len(payday) - 1])[1:-2])
        if ts - payday >= 2592000:
            fam = 'SELECT full_name from users WHERE is_deleted=0 AND NOT LOCATE ("office",login) AND NOT LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(fam, ls)
            fam = mycursor.fetchall()
            fam = str(fam)[3:-5]
            adress = 'SELECT actual_address from users WHERE is_deleted=0 AND NOT LOCATE ("office",login) AND NOT LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(adress, ls)
            adress = mycursor.fetchall()
            adress = str(adress)[3:-5]
            kv = 'SELECT flat_number FROM users WHERE is_deleted=0 AND NOT LOCATE ("office",login) AND NOT LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(kv, ls)
            kv = mycursor.fetchall()
            kv = str(kv)[3:-5]
            bild = 'SELECT building FROM users WHERE is_deleted=0 AND NOT LOCATE ("office",login) AND NOT LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(bild, ls)
            bild = mycursor.fetchall()
            bild = str(bild)[3:-5]
            ost = 'SELECT balance_on_set FROM invoices WHERE is_deleted=0  AND account_id=%s'.format(ls)
            mycursor.execute(ost, ls)
            ost = mycursor.fetchall()
            if len(ost) != 0:
                ost = round(float(str((ost)[-1])[1:-2]), 2)
            else: continue
            if fam.find('РОУТ') != -1:
                fam = fam[6:]
            if bild != '':
                bild = '-' + bild
            print('Фамилия', fam)
            print('Лицевой', ls)
            dts = datetime.today() + timedelta(days=1)
            neopl = datetime.today() + timedelta(days=14)
            s_data = datetime.today()
            ost_data = datetime.today()
            serv = 'SELECT service_id FROM service_links WHERE is_deleted=0 AND account_id=%s'.format(ls)
            mycursor.execute(serv, ls)
            serv = mycursor.fetchall()
            serv = list(serv)
            if len(serv) == 2:
                serv1 = serv[0]
                serv2 = serv[1]
                cena1 = 'SELECT cost FROM periodic_services_data WHERE is_deleted=0 AND id=%s'.format(serv1)
                mycursor.execute(cena1, serv1)
                cena1 = mycursor.fetchall()
                cena1 = round(float(str(cena1)[2:-4]), 2)
                cena2 = 'SELECT cost FROM periodic_services_data WHERE is_deleted=0 AND id=%s'.format(serv2)
                mycursor.execute(cena2, serv2)
                cena2 = mycursor.fetchall()
                cena2 = round(float(str(cena2)[2:-4]), 2)
                print('Цена 1 ', cena1)
                print('Цена 2 ', cena2)
                print('Остаток', ost)
                slink1 = 'SELECT slink_id FROM discount_transactions_all WHERE NOT slink_id=0 AND account_id=%s'.format(
                    ls)
                mycursor.execute(slink1, ls)
                slink1 = mycursor.fetchall()
                slink1 = int(str(slink1[2])[1:-2])
                print('Slink 1 ', slink1)
                slink2 = slink1 + 1
                print('Slink 2 ', slink2)
                stoim1 = 'SELECT discounted FROM dtagg_periodic WHERE slink_id=%s'.format(slink1)
                mycursor.execute(stoim1, slink1)
                stoim1 = mycursor.fetchall()
                stoim1 = round(float(str(stoim1)[2:-4]), 2)
                stoim2 = 'SELECT discounted FROM dtagg_periodic WHERE slink_id=%s'.format(slink2)
                mycursor.execute(stoim2, slink2)
                stoim2 = mycursor.fetchall()
                stoim2 = round(float(str(stoim2)[2:-4]), 2)
                print('Стоимость 1 ', stoim1)
                print('Стоимость 2 ', stoim2)
                stoim = stoim1 + stoim2
                if ost > 0:
                    itogo = round(stoim - ost, 2)
                else:
                    itogo = round(abs(ost) + stoim, 2)
                print()
                doc = DocxTemplate("word2.docx")
                context = {'fam': fam,
                           'data': dts.strftime('%d.%m.%Y'),
                           'neopl': neopl.strftime('%d.%m.%Y'),
                           's_data': s_data.strftime('%d.%m.%Y'),
                           'ost_data': ost_data.strftime('01.%m.%Y'),
                           'adress': adress,
                           'ls': ls,
                           'kv': kv,
                           'bild': bild,
                           'cena1': cena1,
                           'cena2': cena2,
                           'ost': ost,
                           'stoim': stoim,
                           'stoim1': stoim1,
                           'stoim2': stoim2,
                           'itogo': itogo}
                doc.render(context)
                name = str(ls) + '.docx'
                doc.save(name)
                files.append(name)
            elif len(serv) == 1:
                cena = 'SELECT cost FROM periodic_services_data WHERE is_deleted=0 AND id=%s'.format(serv)
                mycursor.execute(cena, serv)
                cena = mycursor.fetchall()
                cena = round(float(str(cena)[2:-4]), 2)
                print('Цена ', cena)
                print('Остаток', ost)
                slink = 'SELECT slink_id FROM discount_transactions_all WHERE NOT slink_id=0 AND account_id=%s'.format(
                    ls)
                mycursor.execute(slink, ls)
                slink = mycursor.fetchall()
                slink = int(str(slink[2])[1:-2])
                print('Slink ', slink)
                stoim = 'SELECT discounted FROM dtagg_periodic WHERE slink_id=%s'.format(slink)
                mycursor.execute(stoim, slink)
                stoim = mycursor.fetchall()
                stoim = round(float(str(stoim)[2:-4]), 2)
                print('Стоимость ', stoim)
                if ost > 0:
                    itogo = round(stoim - ost, 2)
                else:
                    itogo = round(abs(ost) + stoim, 2)
                print()
                doc = DocxTemplate("word1.docx")
                context = {'fam': fam,
                           'data': dts.strftime('%d.%m.%Y'),
                           'neopl': neopl.strftime('%d.%m.%Y'),
                           's_data': s_data.strftime('%d.%m.%Y'),
                           'ost_data': ost_data.strftime('01.%m.%Y'),
                           'adress': adress,
                           'ls': ls,
                           'kv': kv,
                           'bild': bild,
                           'cena': cena,
                           'ost': ost,
                           'stoim': stoim,
                           'itogo': itogo}
                doc.render(context)
                name = str(ls) + '.docx'
                doc.save(name)
                files.append(name)
ids = 1
print()
print('ТВ')
print()
while ids <= 3000:
    mycursor = mydb.cursor()
    ls = 'SELECT basic_account FROM users WHERE is_deleted=0 AND LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND id=%s'.format(
        ids)
    mycursor.execute(ls, ids)
    ls = mycursor.fetchall()
    ls = (str(ls)[2:-4])
    if ls != '':
        ls = int(ls)
        lstv.append(ls)
    ids = ids + 1
for ls in lstv:
    bal = 'SELECT balance FROM accounts WHERE is_deleted=0 AND id=%s'.format(ls)
    mycursor.execute(bal, ls)
    bal = mycursor.fetchall()
    bal = round(float(str(bal)[2:-4]), 2)
    if bal < -30:
        payday = 'SELECT actual_date FROM payment_transactions WHERE account_id=%s'
        mycursor.execute(payday,ls)
        payday = mycursor.fetchall()
        if len(payday) > 0:
            payday = int(str(payday[len(payday) - 1])[1:-2])
        payday = 0
        if ts - payday >= 2592000:
            fam = 'SELECT full_name from users WHERE is_deleted=0 AND LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(fam, ls)
            fam = mycursor.fetchall()
            fam = str(fam)[3:-5]
            adress = 'SELECT actual_address from users WHERE is_deleted=0 AND LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(adress, ls)
            adress = mycursor.fetchall()
            adress = str(adress)[3:-5]
            kv = 'SELECT flat_number FROM users WHERE is_deleted=0  AND LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(kv, ls)
            kv = mycursor.fetchall()
            kv = str(kv)[3:-5]
            bild = 'SELECT building FROM users WHERE is_deleted=0 AND  LOCATE ("tv-",login) AND NOT LOCATE ("РАСТ",full_name) AND NOT LOCATE ("БЛОК",full_name)  AND NOT LOCATE ("ОТКЛ",full_name)  AND NOT LOCATE ("ПАУЗ",full_name) AND basic_account=%s'.format(
                ls)
            mycursor.execute(bild, ls)
            bild = mycursor.fetchall()
            bild = str(bild)[3:-5]
            ost = 'SELECT balance_on_set FROM invoices WHERE is_deleted=0  AND account_id=%s'.format(ls)
            mycursor.execute(ost, ls)
            ost = mycursor.fetchall()
            if len(ost) != 0:
                ost = round(float(str((ost)[-1])[1:-2]), 2)
                if fam.find('РОУТ') != -1:
                    fam = fam[6:]
                if bild != '':
                    bild = '-' + bild
                print('Фамилия', fam)
                print('Лицевой', ls)
                serv = 'SELECT service_id FROM service_links WHERE is_deleted=0 AND account_id=%s'.format(ls)
                mycursor.execute(serv, ls)
                serv = mycursor.fetchall()
                serv = list(serv)
                cena = 'SELECT cost FROM periodic_services_data WHERE is_deleted=0 AND id=%s'.format(serv)
                mycursor.execute(cena, serv)
                cena = mycursor.fetchall()
                cena = round(float(str(cena)[2:-4]), 2)
                print('Цена ', cena)
                print('Остаток', ost)
                slink = 'SELECT slink_id FROM discount_transactions_all WHERE NOT slink_id=0 AND account_id=%s'.format(
                    ls)
                mycursor.execute(slink, ls)
                slink = mycursor.fetchall()
                slink = int(str(slink[0])[1:-2])
                print('Slink ', slink)
                stoim = 'SELECT discounted FROM dtagg_periodic WHERE slink_id=%s'.format(slink)
                mycursor.execute(stoim, slink)
                stoim = mycursor.fetchall()
                stoim = round(float(str(stoim)[2:-4]), 2)
                print('Стоимость ', stoim)
                if ost > 0:
                    itogo = round(stoim - ost, 2)
                else:
                    itogo = round(abs(ost) + stoim, 2)
                print()
                doc = DocxTemplate("word3.docx")
                context = {'fam': fam,
                           'data': dts.strftime('%d.%m.%Y'),
                           'neopl': neopl.strftime('%d.%m.%Y'),
                           's_data': s_data.strftime('%d.%m.%Y'),
                           'ost_data': ost_data.strftime('01.%m.%Y'),
                           'adress': adress,
                           'ls': ls,
                           'kv': kv,
                           'bild': bild,
                           'cena': cena,
                           'ost': ost,
                           'stoim': stoim,
                           'itogo': itogo}
                doc.render(context)
                name = 'tv-' + str(ls) + '.docx'
                doc.save(name)
                files.append(name)
for name in files:
    print(name)
print('Создано ' + str(len(files)) + ' оповещений')
if len(files) > 0:
    dt = dt.strftime("%d.%m.%Y")
    composed = str(dt) + ".docx"
    result = Document(files[0])
    result.add_paragraph()
    composer = Composer(result)
    for i in range(1, len(files)):
        doc = Document(files[i])
        if i != len(files) - 1:
            doc.add_paragraph()
        composer.append(doc)
    composer.save(composed)
    for name in files:
        os.remove(name)
else:
    print('Нет должников')
mydb.commit()
mydb.close()
